from docx import Document
from docx.shared import Inches

doc = Document()
doc.add_heading("마! 이게 파이썬이다")
doc.add_paragraph("탈출은 지능순")


doc.save('pdoc.docx')
